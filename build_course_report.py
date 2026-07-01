from __future__ import annotations

import importlib.util
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT_DOCX = ROOT / "cpr_course_design_report.docx"
ASSET_DIR = ROOT / "report_assets"
ARCH_PNG = ASSET_DIR / "system_architecture.png"

TABLE_GEOMETRY_PATH = Path(
    r"C:\Users\31213\.codex\plugins\cache\openai-primary-runtime\documents"
    r"\26.601.10930\skills\documents\scripts\table_geometry.py"
)


TOC_ENTRIES = [
    ("摘要", "1"),
    ("目录", "2"),
    ("1 研究背景及意义", "3"),
    ("2 研究方法", "4"),
    ("3 研究结果", "7"),
    ("4 各组成员分工", "9"),
    ("5 总结与展望", "10"),
]


def load_table_geometry():
    spec = importlib.util.spec_from_file_location("table_geometry", TABLE_GEOMETRY_PATH)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module.apply_table_geometry


apply_table_geometry = load_table_geometry()


def set_run_font(run, east_asia: str, ascii_font: str, size_pt: float, bold: bool = False):
    run.font.name = ascii_font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def set_paragraph_border_bottom(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFBFBF")
    p_bdr.append(bottom)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def style_body_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(24)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(20)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)

    for run in paragraph.runs:
        set_run_font(run, "宋体", "Times New Roman", 12)


def add_body_paragraph(doc: Document, text: str, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, "宋体", "Times New Roman", 12)
    style_body_paragraph(p, align=align)
    return p


def add_title_like(doc: Document, text: str, size_pt: float, bold: bool = True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, "黑体", "Times New Roman", size_pt, bold=bold)
    return p


def add_heading(doc: Document, text: str, level: int):
    style_name = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
    p = doc.add_paragraph(style=style_name)
    p.clear()
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, "黑体", "Times New Roman", 15, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt = p.paragraph_format
        fmt.first_line_indent = Pt(0)
        fmt.space_before = Pt(30)
        fmt.space_after = Pt(18)
    elif level == 2:
        set_run_font(run, "黑体", "Times New Roman", 14, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt = p.paragraph_format
        fmt.first_line_indent = Pt(0)
        fmt.space_before = Pt(18)
        fmt.space_after = Pt(12)
    else:
        set_run_font(run, "黑体", "Times New Roman", 14, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt = p.paragraph_format
        fmt.first_line_indent = Pt(0)
        fmt.space_before = Pt(12)
        fmt.space_after = Pt(6)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(20)
    return p


def add_table_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.space_before = Pt(12)
    fmt.space_after = Pt(6)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(20)
    run = p.add_run(text)
    set_run_font(run, "宋体", "Times New Roman", 12, bold=True)
    return p


def add_figure_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(12)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(20)
    run = p.add_run(text)
    set_run_font(run, "宋体", "Times New Roman", 12)
    return p


def set_table_text(cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    fmt = p.paragraph_format
    fmt.first_line_indent = Pt(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(18)
    run = p.add_run(text)
    set_run_font(run, "宋体", "Times New Roman", 10.5, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def create_architecture_figure(path: Path):
    from PIL import Image, ImageDraw, ImageFont

    path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1800, 900), "white")
    draw = ImageDraw.Draw(img)

    font_path = r"C:\Windows\Fonts\msyh.ttc"
    title_font = ImageFont.truetype(font_path, 42)
    text_font = ImageFont.truetype(font_path, 30)
    small_font = ImageFont.truetype(font_path, 24)

    draw.text((620, 40), "智能CPR手势评估系统总体流程", fill="#222222", font=title_font)

    boxes = [
        ((80, 220, 360, 360), "#EAF2FF", "摄像头采集", "OpenCV读取视频帧"),
        ((430, 220, 760, 360), "#EAF7EA", "人体检测", "YOLOv8n识别人框"),
        ((830, 220, 1160, 360), "#FFF5E6", "姿态估计", "YOLOv8-pose输出17点"),
        ((1230, 220, 1560, 360), "#FBEAEA", "结果融合", "合并bbox与关键点"),
        ((430, 520, 760, 700), "#F4ECFF", "CPR分析", "角度 / BPM / 位置 / 回弹 / 深度"),
        ((830, 520, 1160, 700), "#EAF2FF", "可视化叠加", "骨架、标记与文字提示"),
        ((1230, 520, 1560, 700), "#EAF7EA", "图形界面", "PyQt5实时显示与控制"),
    ]

    for x1, y1, x2, y2 in [b[0] for b in boxes]:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=24, outline="#4A4A4A", width=3)
    for (x1, y1, x2, y2), fill, title, subtitle in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=24, fill=fill, outline="#4A4A4A", width=3)
        draw.text((x1 + 30, y1 + 36), title, fill="#1F1F1F", font=text_font)
        draw.text((x1 + 30, y1 + 98), subtitle, fill="#444444", font=small_font)

    arrows = [
        ((360, 290), (430, 290)),
        ((760, 290), (830, 290)),
        ((1160, 290), (1230, 290)),
        ((595, 360), (595, 520)),
        ((995, 360), (995, 520)),
        ((1395, 360), (1395, 520)),
        ((760, 610), (830, 610)),
        ((1160, 610), (1230, 610)),
    ]

    for start, end in arrows:
        draw.line([start, end], fill="#5B6B8C", width=6)
        dx = end[0] - start[0]
        dy = end[1] - start[1]
        if abs(dx) > abs(dy):
            if dx > 0:
                pts = [(end[0], end[1]), (end[0] - 24, end[1] - 12), (end[0] - 24, end[1] + 12)]
            else:
                pts = [(end[0], end[1]), (end[0] + 24, end[1] - 12), (end[0] + 24, end[1] + 12)]
        else:
            if dy > 0:
                pts = [(end[0], end[1]), (end[0] - 12, end[1] - 24), (end[0] + 12, end[1] - 24)]
            else:
                pts = [(end[0], end[1]), (end[0] - 12, end[1] + 24), (end[0] + 12, end[1] + 24)]
        draw.polygon(pts, fill="#5B6B8C")

    draw.text((120, 760), "系统采用轻量级YOLO模型与桌面可视化界面，支持实时评估与后续标定扩展。", fill="#4A4A4A", font=small_font)
    img.save(path)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.left_margin = Mm(31.75)
    section.right_margin = Mm(31.75)
    section.header_distance = Mm(15)
    section.footer_distance = Mm(17.5)
    section.start_type = WD_SECTION.NEW_PAGE

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style.paragraph_format.line_spacing = Pt(20)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.first_line_indent = Pt(24)

    for style_name, east_asia, size_pt in [
        ("Heading 1", "黑体", 15),
        ("Heading 2", "黑体", 14),
        ("Heading 3", "黑体", 14),
    ]:
        s = doc.styles[style_name]
        s.font.name = "Times New Roman"
        s.font.size = Pt(size_pt)
        s.font.bold = True
        s._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        s._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def add_cover(doc: Document):
    for _ in range(6):
        doc.add_paragraph()
    add_title_like(doc, "医学仪器原理与设计", 22)
    doc.add_paragraph()
    add_title_like(doc, "课程设计报告", 22)
    for _ in range(5):
        doc.add_paragraph()

    items = [
        "设计题目：智能CPR手势评估系统",
        "学科专业：待填写",
        "作者姓名：待填写",
        "指导教师：王慧泉",
        "完成日期：2026-06-15",
    ]
    for text in items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt = p.paragraph_format
        fmt.first_line_indent = Pt(0)
        fmt.space_before = Pt(12)
        fmt.space_after = Pt(0)
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = Pt(20)
        run = p.add_run(text)
        set_run_font(run, "宋体", "Times New Roman", 14)

    doc.add_page_break()


def add_abstract(doc: Document):
    add_heading(doc, "摘要", 1)
    abstract = (
        "心肺复苏是临床急救和公共急救培训中的关键技能，但传统教学与考核方式较多依赖"
        "教师现场观察，存在主观性强、记录粒度不足和难以形成连续量化反馈等问题。围绕这一"
        "痛点，本文结合现有项目代码，设计并实现了一套基于计算机视觉的智能CPR手势评估系"
        "统。系统以普通摄像头为输入，采用YOLOv8n完成人体检测，采用YOLOv8n-pose提取17个"
        "人体关键点，并在此基础上构建按压动作分析模块。通过对肘关节角度、按压频率、手部"
        "位置、回弹状态及肩腕相对距离变化等信息进行综合判定，系统能够对施救者的按压质量"
        "进行实时评估，并在PyQt5界面中同步显示骨架叠加、统计结果与运行状态。项目目前已完"
        "成模型加载、摄像头采集、多人角色选择、标定、可视化和基础测试等功能，具备课堂演"
        "示与训练辅助价值。结果表明，该方案能够在轻量级桌面平台上完成CPR动作的实时分析，"
        "为急救技能教学提供一种低成本、可扩展的智能化实现路径。但由于系统使用单目摄像头，"
        "对按压深度的绝对测量能力仍然有限，后续仍需结合深度传感器或多视角方法进一步提升"
        "精度与鲁棒性。"
    )
    add_body_paragraph(doc, abstract)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.first_line_indent = Pt(0)
    fmt.space_before = Pt(12)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(20)
    key_label = p.add_run("关键词：")
    set_run_font(key_label, "黑体", "Times New Roman", 12, bold=True)
    key_text = p.add_run("心肺复苏；姿态估计；YOLOv8；PyQt5；实时评估")
    set_run_font(key_text, "宋体", "Times New Roman", 12)
    doc.add_page_break()


def add_toc(doc: Document):
    add_heading(doc, "目录", 1)
    for title, page in TOC_ENTRIES:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt = p.paragraph_format
        fmt.first_line_indent = Pt(0)
        fmt.left_indent = Pt(24 if title.startswith("1 ") or title.startswith("2 ") or title.startswith("3 ") or title.startswith("4 ") or title.startswith("5 ") else 0)
        fmt.right_indent = Pt(0)
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = Pt(20)

        tab_stops = fmt.tab_stops
        tab_stops.add_tab_stop(Cm(11.8))
        run = p.add_run(f"{title}\t{page}")
        set_run_font(run, "宋体", "Times New Roman", 12)
    doc.add_page_break()


def add_module_table(doc: Document):
    add_table_title(doc, "表2-1 系统主要功能模块")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["模块名称", "核心文件", "主要作用"]
    for idx, text in enumerate(headers):
        set_table_text(table.rows[0].cells[idx], text, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "E8EEF5")

    rows = [
        ("摄像头采集", "camera/camera_thread.py", "使用QThread持续采集视频流并向界面发送结果。"),
        ("目标检测", "detector/object_detector.py", "调用YOLOv8n检测画面中的人体目标框。"),
        ("姿态估计", "detector/pose_estimator.py", "调用YOLOv8n-pose输出17个关键点坐标。"),
        ("动作分析", "analyzer/cpr_analyzer.py", "计算角度、频率、手位、回弹与相对深度。"),
        ("图形界面", "ui/main_window.py 等", "显示视频、统计信息、状态提示和交互控件。"),
    ]
    for row in rows:
        cells = table.add_row().cells
        set_table_text(cells[0], row[0], align=WD_ALIGN_PARAGRAPH.CENTER)
        set_table_text(cells[1], row[1], align=WD_ALIGN_PARAGRAPH.LEFT)
        set_table_text(cells[2], row[2], align=WD_ALIGN_PARAGRAPH.LEFT)

    apply_table_geometry(table, [1800, 2600, 4960], table_width_dxa=9360, indent_dxa=120)


def add_metrics_table(doc: Document):
    add_table_title(doc, "表2-2 CPR关键评价指标与判定依据")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["评价指标", "判定依据", "目标范围或逻辑"]
    for idx, text in enumerate(headers):
        set_table_text(table.rows[0].cells[idx], text, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "E8EEF5")

    rows = [
        ("肘关节角度", "肩-肘-腕三点夹角", "角度不小于160°时认为手臂基本伸直。"),
        ("按压频率", "腕部Y坐标变化的时间间隔", "在5 s滑动窗口中换算为100~120 bpm。"),
        ("手部位置", "腕部相对人体边界框的位置", "位于身体中心附近或标定区域内。"),
        ("回弹状态", "腕部是否回到基线附近", "回到基线±2%范围内视为回弹充分。"),
        ("相对深度", "肩腕距离相对基线的变化", "变化幅度达到阈值时提示按压有效。"),
    ]
    for row in rows:
        cells = table.add_row().cells
        set_table_text(cells[0], row[0], align=WD_ALIGN_PARAGRAPH.CENTER)
        set_table_text(cells[1], row[1], align=WD_ALIGN_PARAGRAPH.LEFT)
        set_table_text(cells[2], row[2], align=WD_ALIGN_PARAGRAPH.LEFT)

    apply_table_geometry(table, [1500, 2700, 5160], table_width_dxa=9360, indent_dxa=120)


def add_env_table(doc: Document):
    add_table_title(doc, "表3-1 系统开发与验证环境")
    table = doc.add_table(rows=1, cols=3)
    headers = ["组件", "版本", "用途"]
    for idx, text in enumerate(headers):
        set_table_text(table.rows[0].cells[idx], text, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "E8EEF5")

    rows = [
        ("Python", "3.12", "作为项目的主要开发语言和运行环境。"),
        ("PyTorch", "2.6.0+cu126", "承担深度学习模型加载与推理任务。"),
        ("Ultralytics", "8.4.47", "封装YOLOv8n与YOLOv8n-pose模型。"),
        ("OpenCV", "4.13.0", "负责摄像头采集、图像转换与图形绘制。"),
        ("PyQt5", "5.15.10", "构建桌面图形界面与交互逻辑。"),
        ("NumPy", "2.1.2", "完成数组与数值计算支撑。"),
    ]
    for row in rows:
        cells = table.add_row().cells
        set_table_text(cells[0], row[0], align=WD_ALIGN_PARAGRAPH.CENTER)
        set_table_text(cells[1], row[1], align=WD_ALIGN_PARAGRAPH.CENTER)
        set_table_text(cells[2], row[2], align=WD_ALIGN_PARAGRAPH.LEFT)

    apply_table_geometry(table, [1800, 1800, 5760], table_width_dxa=9360, indent_dxa=120)


def add_validation_table(doc: Document):
    add_table_title(doc, "表3-2 已完成功能与验证情况")
    table = doc.add_table(rows=1, cols=3)
    headers = ["验证项目", "结果", "说明"]
    for idx, text in enumerate(headers):
        set_table_text(table.rows[0].cells[idx], text, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "E8EEF5")

    rows = [
        ("模块导入", "通过", "各子模块能够在Python环境中正常导入。"),
        ("检测模型加载", "通过", "yolov8n.pt 文件存在且可被程序调用。"),
        ("姿态模型加载", "通过", "yolov8n-pose.pt 文件存在且可被程序调用。"),
        ("主界面初始化", "通过", "PyQt5窗口、状态栏和参数面板均可创建。"),
        ("分析器单元校验", "通过", "角度计算、空输入与模拟按压逻辑工作正常。"),
        ("多模块联调", "通过", "摄像头线程、检测器、分析器与UI初始化顺畅。"),
    ]
    for row in rows:
        cells = table.add_row().cells
        set_table_text(cells[0], row[0], align=WD_ALIGN_PARAGRAPH.CENTER)
        set_table_text(cells[1], row[1], align=WD_ALIGN_PARAGRAPH.CENTER)
        set_table_text(cells[2], row[2], align=WD_ALIGN_PARAGRAPH.LEFT)

    apply_table_geometry(table, [2200, 1200, 5960], table_width_dxa=9360, indent_dxa=120)


def add_member_table(doc: Document):
    table = doc.add_table(rows=1, cols=2)
    headers = ["成员", "分工"]
    for idx, text in enumerate(headers):
        set_table_text(table.rows[0].cells[idx], text, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "E8EEF5")

    rows = [
        ("成员1（待填写）", "需求分析、系统设计、代码实现、功能测试、报告撰写。"),
        ("成员2（可选）", "若为小组课设，可按实际情况补充界面设计、实验测试或文档整理分工。"),
    ]
    for row in rows:
        cells = table.add_row().cells
        set_table_text(cells[0], row[0], align=WD_ALIGN_PARAGRAPH.CENTER)
        set_table_text(cells[1], row[1], align=WD_ALIGN_PARAGRAPH.LEFT)

    apply_table_geometry(table, [2200, 7160], table_width_dxa=9360, indent_dxa=120)


def build_document():
    create_architecture_figure(ARCH_PNG)
    doc = Document()
    configure_document(doc)

    add_cover(doc)
    add_abstract(doc)
    add_toc(doc)

    add_heading(doc, "1 研究背景及意义", 1)
    add_body_paragraph(
        doc,
        "心肺复苏训练强调按压深度、频率、回弹和手位等核心动作质量，但在常规课堂教学中，"
        "教师往往只能通过肉眼观察和口头纠正来完成评估。这种方式虽然直观，却很难形成稳定、"
        "连续且可回溯的数据记录，尤其在多人训练场景下，教师难以及时兼顾每一名学生的动作细节。"
    )
    add_body_paragraph(
        doc,
        "随着计算机视觉技术与轻量级深度学习模型的发展，普通摄像头加姿态估计的方法已经能够"
        "在不增加复杂传感器成本的前提下提取人体关键点信息，为急救动作分析提供新的技术手段。"
        "将其引入CPR教学，既可以提高课堂反馈效率，也便于后续进行量化统计与训练过程复盘。"
    )
    add_body_paragraph(
        doc,
        "本项目面向《医学仪器原理与设计》课程设计场景，目标是在现有桌面设备条件下构建一套"
        "可运行、可演示、可扩展的智能CPR手势评估系统。系统不追求替代专业急救训练仪器，而是"
        "重点验证视觉评估方案在教学辅助、动作提示和系统集成方面的可行性。"
    )
    add_heading(doc, "1.1 课题研究价值", 2)
    add_body_paragraph(
        doc,
        "从教学角度看，该系统能够将抽象的动作规范转化为实时指标提示，帮助学习者更快理解"
        "“按得是否够快、手臂是否伸直、回弹是否充分”等关键问题；从工程实现角度看，本项目覆"
        "盖了图像采集、模型推理、时序分析、桌面界面设计与模块协同，具有较强的综合训练意义。"
    )
    add_body_paragraph(
        doc,
        "此外，该系统还为后续扩展留下了接口。例如，在保持现有界面的前提下，可继续引入更高"
        "精度的深度估计、课堂记录存档、训练成绩导出以及多摄像头协同分析等功能，因此具有一定"
        "的后续开发价值。"
    )

    doc.add_page_break()
    add_heading(doc, "2 研究方法", 1)
    add_heading(doc, "2.1 系统需求与总体方案", 2)
    add_body_paragraph(
        doc,
        "围绕CPR训练辅助这一应用目标，系统需要满足以下基本要求：第一，能够实时采集摄像头"
        "视频并识别画面中的施救者；第二，能够从人体关键点中提取与按压动作密切相关的姿态信息；"
        "第三，能够将分析结果直观显示在界面中，便于课堂演示与即时纠正；第四，整体实现需尽量"
        "轻量，能够在普通CPU环境下完成运行。"
    )
    add_body_paragraph(
        doc,
        "为实现上述要求，项目采用“视频采集 + 目标检测 + 姿态估计 + CPR规则分析 + 图形界面显"
        "示”的总体方案。通过YOLOv8n快速定位人体边界框，再由YOLOv8n-pose输出17个关键点，最后"
        "结合设定阈值完成动作质量判定。"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run()
    run.add_picture(str(ARCH_PNG), width=Cm(15.5))
    add_figure_title(doc, "图2-1 智能CPR手势评估系统总体流程")

    add_heading(doc, "2.2 模块化设计", 2)
    add_body_paragraph(
        doc,
        "项目采用分层模块结构，将摄像头线程、检测模块、姿态模块、分析模块和界面模块解耦，"
        "有利于后续维护与独立测试。核心模块与作用如表2-1所示。"
    )
    add_module_table(doc)
    add_heading(doc, "2.3 人体检测与姿态估计方法", 2)
    add_body_paragraph(
        doc,
        "人体检测模块基于YOLOv8n实现，重点提取COCO数据集中类别为person的目标框。由于YOLOv8n"
        "属于轻量级检测模型，其参数量和推理开销相对较小，适合课程设计中的实时演示场景。"
    )
    add_body_paragraph(
        doc,
        "姿态估计模块使用YOLOv8n-pose，对每个检测到的人体输出17个关键点，包括肩、肘、腕、"
        "髋、膝、踝等部位。系统在多人场景下会融合检测框与姿态结果，并按画面位置进行排序，供"
        "界面侧的施救者/被救者编号选择控件使用。"
    )
    add_body_paragraph(
        doc,
        "为减少关键点抖动带来的误判，系统对关键点序列增加了时序平滑策略：对施救者采用较轻"
        "的平滑系数，以保留动作响应速度；对其他人物采用较强平滑，以提高画面稳定性。这一设计"
        "兼顾了评估灵敏度与显示可读性。"
    )

    add_heading(doc, "2.4 CPR动作评估方法", 2)
    add_body_paragraph(
        doc,
        "CPR分析模块的核心思想是将视觉关键点映射为可解释的训练指标。系统优先选择置信度较高"
        "的一侧上肢，基于肩、肘、腕三点计算肘关节角度，以判断手臂是否伸直；通过腕部Y坐标在时"
        "间轴上的变化识别一次新的按压，并在滑动时间窗口内估算按压频率；通过腕部相对于人体边"
        "界框中心的位置判断手位是否合理；通过腕部是否回到基线附近判断回弹是否充分。"
    )
    add_body_paragraph(
        doc,
        "考虑到项目使用单目摄像头，系统无法直接给出厘米级按压深度，因此采用肩腕距离相对于"
        "标定基线的变化量作为相对深度指标。这种方法虽然不能替代专业压力传感器，但可以有效区"
        "分“存在明显按压”与“动作幅度不足”的情况。"
    )
    add_metrics_table(doc)

    add_heading(doc, "2.5 界面交互与运行流程", 2)
    add_body_paragraph(
        doc,
        "界面使用PyQt5实现，左侧为视频显示区域，右侧为实时参数面板。用户可以完成摄像头选择、"
        "开始/停止评估、统计重置、ROI标定以及施救者与被救者编号设定。界面顶部状态栏用于显示"
        "系统当前运行状态与FPS信息，便于课堂演示时快速观察运行情况。"
    )
    add_body_paragraph(
        doc,
        "在运行流程上，系统首先完成模型与界面初始化；随后由摄像头线程持续读取视频帧；每一帧"
        "依次经过人体检测、姿态估计、人物融合和CPR分析；最后将骨架连线、人物边框和文本结果叠"
        "加到图像上，并同步更新右侧面板中的角度、频率、回弹和统计值。"
    )

    doc.add_page_break()
    add_heading(doc, "3 研究结果", 1)
    add_heading(doc, "3.1 软硬件环境", 2)
    add_body_paragraph(
        doc,
        "本项目采用Python语言开发，核心依赖包括PyTorch、Ultralytics、OpenCV、PyQt5和NumPy。"
        "根据项目工作日志，当前模型文件已经位于本地，可在不联网的前提下完成系统启动与功能演示。"
        "主要开发与验证环境如表3-1所示。"
    )
    add_env_table(doc)
    add_heading(doc, "3.2 已实现功能", 2)
    add_body_paragraph(
        doc,
        "从现有代码实现来看，系统已经具备完整的桌面端运行骨架。其一，主程序能够创建应用窗"
        "口并组装视频区和信息面板；其二，摄像头线程可以独立执行视频采集与模型推理；其三，动"
        "作分析器能够对按压相关关键指标进行判定；其四，界面支持角色选择、标定和状态显示。"
    )
    add_body_paragraph(
        doc,
        "与单一的演示脚本相比，本项目更加突出工程组织性：各功能模块分布于独立目录中，配置"
        "项集中于config.py，模型下载脚本单独维护，便于后续继续扩展实验场景或调整阈值。"
    )
    add_heading(doc, "3.3 验证情况与分析", 2)
    add_body_paragraph(
        doc,
        "结合代码检查结果和现有工作日志，系统当前已完成的主要验证项目如表3-2所示。可以看出，"
        "项目已经完成从模型准备、模块导入到界面联调的基本闭环，具备课程设计展示和后续完善的"
        "工程基础。"
    )
    add_validation_table(doc)
    add_body_paragraph(
        doc,
        "从结果质量看，系统在教学辅助中的优势主要体现在三个方面：第一，评价指标直接对应CPR"
        "训练中的关键动作要求，解释性较强；第二，系统采用轻量化模型和桌面界面，部署成本较低；"
        "第三，项目保留了标定和多人角色选择逻辑，说明作者在实际课堂场景适配上已有考虑。"
    )
    add_heading(doc, "3.4 存在问题与改进方向", 2)
    add_body_paragraph(
        doc,
        "目前系统的主要局限在于按压深度仍为相对估计值，无法替代专业训练设备给出的绝对深度"
        "测量；另外，摄像头视角、遮挡、光照变化以及多人互相干扰都可能影响关键点稳定性。"
    )
    add_body_paragraph(
        doc,
        "后续可从以下方向继续改进：一是引入深度摄像头或双目视觉，提高按压深度估计精度；二"
        "是增加动作数据记录与成绩导出模块，用于训练过程管理；三是结合自动角色分类逻辑，减少"
        "用户手动设置施救者/被救者编号的操作负担。"
    )

    doc.add_page_break()
    add_heading(doc, "4 各组成员分工", 1)
    add_body_paragraph(
        doc,
        "考虑到模板中要求明确分工信息，本文给出可直接替换的分工表。若本次课程设计为个人完成，"
        "可仅保留第一行并将成员姓名替换为本人；若为小组项目，可在现有表格基础上继续增补。"
    )
    add_member_table(doc)

    doc.add_page_break()
    add_heading(doc, "5 总结与展望", 1)
    add_body_paragraph(
        doc,
        "本文依据当前项目代码完成了智能CPR手势评估系统的课程设计文档整理。项目以YOLOv8n和"
        "YOLOv8n-pose为核心视觉模型，结合桌面界面与规则分析模块，实现了对CPR关键动作质量的实"
        "时评估。整体方案在普通摄像头和通用桌面平台上即可运行，体现出较好的课程实践价值。"
    )
    add_body_paragraph(
        doc,
        "从课程设计角度看，该项目不仅关注算法本身，还兼顾了模块划分、界面交互、系统联调和"
        "工程可维护性，较好地体现了“医学仪器原理与设计”课程中软硬件协同、数据处理和用户交互"
        "设计的综合要求。"
    )
    add_body_paragraph(
        doc,
        "后续工作可继续围绕精度提升、场景扩展和教学应用深化三个方向展开，包括引入更稳定的"
        "深度感知方案、构建带标签的CPR训练数据集、加入训练日志与评分记录模块，以及探索面向多"
        "人课堂的批量化训练辅助功能。"
    )

    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_document()
